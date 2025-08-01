# # ingest.py

import os
import re
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.docstore.document import Document as LangchainDocument
from config import EMBEDDING_MODEL_NAME, CHUNK_SIZE, CHUNK_OVERLAP, CLEANED_DIR, PDF_DIR, VECTORSTORE_DIR

# Add File Visibility to Metadata For RBAC purposes - Charmaine
import mysql.connector

DB_CONFIG = {
    'host': 'localhost',
    'user': 'user',
    'password': 'password',
    'database': 'Verztec'
}

def get_visibility_for_file(file_base_name):
    connection = mysql.connector.connect(**DB_CONFIG)
    cursor = connection.cursor(dictionary=True)
    # Get file ID (assuming file_name in DB is like 'filename.pdf')
    # Remove .pdf, .docx, .txt (case-insensitive) from filename for matching
    cursor.execute(
        "SELECT id FROM files WHERE LOWER(REPLACE(REPLACE(REPLACE(REPLACE(REPLACE(filename, '.pdf', ''), '.docx', ''), '.doc', ''), '.txt', ''), '.xlsx', '')) = %s",
        (file_base_name.lower(),)
    )
    file_row = cursor.fetchone()
    if not file_row:
        cursor.close()
        connection.close()
        return {"countries": "UNKNOWN"}

    file_id = file_row["id"]

    # Retrieve visibility from file_visibility table
    cursor.execute("SELECT DISTINCT country, department FROM file_visibility WHERE file_id = %s", (file_id,))
    rows = cursor.fetchall()
    cursor.close()
    connection.close()

    if not rows:
        return {"countries": "NONE"}

    # If any row has country set to ALL, file is visible to everyone
    if any(row["country"] == "ALL" and row["department"] == "ALL" for row in rows):
        return {"countries": "ALL"}

    # Collect unique countries (excluding ALL)
    countries = set()
    for row in rows:
        if row["country"] != "ALL":
            countries.add(row["country"])
    
    # Return space-separated list of countries for vectorstore metadata
    if countries:
        return {"countries": " ".join(sorted(countries))}
    else:
        return {"countries": "RESTRICTED"}


# 🔍 Read DOCX files
def read_docx(file_path):
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

# 🧹 Optional: additional cleaning
def clean_text(text):
    text = re.sub(r"\*|\d+\.", "", text)
    return text.strip()

# 🚀 Main ingestion function
def ingest_documents():
    splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
    embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL_NAME)
    docs = []

    for file in os.listdir(CLEANED_DIR):
        path = os.path.join(CLEANED_DIR, file)

        # 🔄 Load text content
        if file.endswith(".txt"):
            with open(path, encoding="utf-8") as f:
                text = f.read()
            base_name = file.replace(".txt", "")
        elif file.endswith(".docx"):
            text = read_docx(path)
            base_name = file.replace(".docx", "")
        else:
            continue

        text = clean_text(text)
        chunks = splitter.create_documents([text])

        # 🔗 Match with original file
        matched_file = None
        for f in os.listdir(PDF_DIR):
            filename_wo_ext, _ = os.path.splitext(f)
            if base_name.lower() == filename_wo_ext.lower():
                matched_file = f
                break

        source_file = matched_file if matched_file else f"{base_name}.docx"

        # 📄 Use first chunk to intelligently detect doc_type
        first_chunk_text = chunks[0].page_content if chunks else ""
        # 🏷️ Metadata tagging
        def infer_doc_type(base_name, first_chunk_text):
            lowered = first_chunk_text.lower()

            if any(keyword in lowered for keyword in ["quality manual", "quality procedure"]) and "controlled copy" in lowered:
                return "cover_page"
            elif any(term in base_name.lower() for term in ["digital meeting", "online meeting", "virtual meeting"]):
                return "digital"
            elif "etiquette" in base_name.lower() or "physical" in base_name.lower():
                return "physical"
            else:
                return "general"
        doc_type = infer_doc_type(base_name, first_chunk_text)
        
        # if "cover" in base_name.lower():
        #     doc_type = "cover_page"
        # elif any(term in base_name.lower() for term in ["digital meeting", "online meeting", "virtual meeting"]):
        #     doc_type = "digital"
        # elif "etiquette" in base_name.lower() or "physical" in base_name.lower():
        #     doc_type = "physical"
        # else:
        #     doc_type = "general"

        # get Visibility info of file for RBAC (Charmaine)
        visibility = get_visibility_for_file(base_name)

        # 📎 Attach metadata to each chunk
        for chunk in chunks:
            chunk.metadata["source"] = source_file
            chunk.metadata["title"] = base_name.replace("_", " ").lower().strip()
            chunk.metadata["doc_type"] = doc_type
            chunk.metadata["countries"] = visibility["countries"]  # RBAC (Charmaine)
            docs.append(chunk)

    # 💾 Save to FAISS
    vectorstore = FAISS.from_documents(docs, embeddings)
    vectorstore.save_local(VECTORSTORE_DIR)
    print("✅ Ingestion complete.")

if __name__ == "__main__":
    ingest_documents()


# import os
# import re
# from docx import Document
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain_community.vectorstores import FAISS
# # from langchain.embeddings import HuggingFaceEmbeddings
# from langchain_community.embeddings import HuggingFaceEmbeddings
# from langchain.docstore.document import Document as LangchainDocument
# from config import EMBEDDING_MODEL_NAME, CHUNK_SIZE, CHUNK_OVERLAP, CLEANED_DIR, PDF_DIR, VECTORSTORE_DIR

# def read_docx(file_path):
#     doc = Document(file_path)
#     return "\n".join([para.text for para in doc.paragraphs])

# def clean_text(text):
#     text = re.sub(r"\*|\d+\.", "", text)
#     return text.strip()

# def ingest_documents():
#     splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
#     embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL_NAME)
#     docs = []

#     for file in os.listdir(CLEANED_DIR):
#         path = os.path.join(CLEANED_DIR, file)

#         if file.endswith(".txt"):
#             with open(path, encoding="utf-8") as f:
#                 text = f.read()
#             base_name = file.replace(".txt", "")
#         elif file.endswith(".docx"):
#             text = read_docx(path)
#             base_name = file.replace(".docx", "")
#         else:
#             continue

#         text = clean_text(text)
#         chunks = splitter.create_documents([text])

#         # Find a matching PDF or DOCX in the PDF_DIR folder
#         # matched_file = None
#         # for f in os.listdir(PDF_DIR):
#         #     if base_name in f:
#         #         matched_file = f
#         #         break

#         matched_file = None
#         for f in os.listdir(PDF_DIR):
#             filename_wo_ext, _ = os.path.splitext(f)
#             if base_name.lower() == filename_wo_ext.lower():
#                 matched_file = f
#                 break


#         source_file = matched_file if matched_file else f"{base_name}.docx"

#         for chunk in chunks:
#             chunk.metadata["source"] = source_file
#             chunk.metadata["title"] = base_name.replace("_", " ").lower().strip()
#             docs.append(chunk)

#     vectorstore = FAISS.from_documents(docs, embeddings)
#     vectorstore.save_local(VECTORSTORE_DIR)
#     print("✅ Ingestion complete.")

# if __name__ == "__main__":
#     ingest_documents()


